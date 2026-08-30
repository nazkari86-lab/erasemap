#!/usr/bin/env python3
# ruff: noqa: E501, RUF001
"""Build the public, non-sensitive EraSeMap submission documents."""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[2]
PAPER_DIR = ROOT / "competition" / "paper"
OUT = ROOT / "competition" / "submission"
sys.path.insert(0, str(PAPER_DIR))

from build_papers import (  # noqa: E402
    AUTHOR_RU,
    BLUE,
    DARK,
    DIRECTION_RU,
    MID_BLUE,
    MUTED,
    SCHOOL_RU,
    SECTION_RU,
    SUPERVISOR_RU,
    add_inline,
    configure_document,
    set_run_font,
)

TITLE = "EraSeMap: проверяемое удаление персональных данных из распределённых систем и моделей машинного обучения"

ABSTRACT_RU = (
    "Команда удаления основной записи не гарантирует, что персональные данные перестали использоваться. "
    "Копии могут сохраняться в кэше, индексе, реплике, экспорте или резервной копии; производные — "
    "в биометрическом шаблоне и векторе; влияние — в обученной модели. Я разработал EraSeMap — один "
    "алгоритм из трёх стадий. FIND находит известные и ограниченно неочевидные пути. ERASE выбирает "
    "минимальный достаточный набор физических действий и удаление влияния из модели (machine unlearning). "
    "PROVE повторяет будущие операции и разрешает сертификат только после прохождения всех обязательных "
    "проверок. В заранее зафиксированной проверке переноса на 60 случаях EraSeMap дал 0 ложных COMPLETE "
    "против 5 у полного аудита узлов и 45 у обычного статуса сервиса. Активный поиск ограниченного "
    "графа восстановления потребовал 7 проб против 13 у случайной стратегии и 49 у полного перебора. "
    "В 20 парных опытах с реальными локальными процессами выбранный план был в 17,64 раза быстрее полной "
    "пересборки и записал на 94,62% меньше байтов. Временная проверка обнаружила 30 из 30 скрытых рисков "
    "и после контрольных действий дала 0 из 30 возвратов. Быстрые методы удаления влияния из Qwen2.5-1.5B "
    "не прошли все заранее заданные пороги, поэтому полное переобучение оставлено безопасным запасным "
    "вариантом. Результаты ограничены зарегистрированными топологиями и переходами; промышленное внедрение "
    "и независимая скрытая проверка не заявляются."
)

ABSTRACT_EN = (
    "Deleting a primary record does not prove that personal data can no longer be used. Copies may remain "
    "in caches, indexes, replicas, exports, or backups; derivatives may remain as biometric templates and "
    "vectors; learned influence may remain in a trained model. I developed EraSeMap as one algorithm "
    "with three stages. FIND locates known and bounded, non-obvious paths. ERASE selects a minimum sufficient "
    "set of physical actions and removes learned influence from the model (machine unlearning). PROVE repeats "
    "future operations and permits a certificate only after every required check passes. In a pre-registered "
    "60-case transfer, EraSeMap produced 0 false COMPLETE decisions versus 5 for a full node audit and 45 "
    "for the service's ordinary status. Active recovery-graph search used 7 probes versus 13 for random testing "
    "and 49 for exhaustive testing. In 20 paired real-process trials, the selected plan was 17.64 times faster "
    "and wrote 94.62% fewer bytes than rebuilding everything. Temporal "
    "verification detected 30 of 30 latent risks and produced 0 of 30 post-control recurrences. Fast unlearning "
    "candidates on Qwen2.5-1.5B failed all required thresholds together, so full retraining from scratch remains "
    "the safe fallback. Results are limited to registered topology and listed transitions; production deployment "
    "and an independent hidden result are not claimed."
)


def word_count(text: str) -> int:
    return len(text.replace("—", " ").split())


def add_document_title(doc: Document, label: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(label)
    set_run_font(r, name="PT Sans", size=11, bold=True, color=MID_BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(TITLE)
    set_run_font(r, name="PT Sans", size=18, bold=True, color=BLUE)


def add_identity(doc: Document) -> None:
    entries = [
        f"Автор: {AUTHOR_RU}",
        f"Организация: {SCHOOL_RU}",
        DIRECTION_RU,
        SECTION_RU,
        f"Научный руководитель: {SUPERVISOR_RU}",
        "Алматы, 2026",
    ]
    for entry in entries:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(entry)
        set_run_font(r, name="PT Sans", size=9.8, color=MUTED)
    doc.add_paragraph()


def add_section(doc: Document, heading: str, text: str) -> None:
    doc.add_paragraph(heading, style="Heading 1")
    p = doc.add_paragraph()
    p.paragraph_format.widow_control = True
    add_inline(p, text)


def build_theses(path: Path) -> None:
    doc = Document()
    configure_document(doc, True)
    add_document_title(doc, "ТЕЗИСЫ НАУЧНО-ИССЛЕДОВАТЕЛЬСКОЙ РАБОТЫ")
    add_identity(doc)
    add_section(
        doc,
        "Проблема",
        "Ответ сервиса DELETE подтверждает выполнение одной команды, но не исчезновение копий, "
        "биометрических шаблонов, поисковых векторов, резервных копий и влияния в обученной модели. "
        "Кроме того, удалённые данные могут появиться снова после восстановления или синхронизации.",
    )
    add_section(
        doc,
        "Цель и гипотеза",
        "Цель — разработать один понятный и проверяемый алгоритм удаления персональных данных, который учитывает "
        "физические копии, производные, влияние в модели и будущие пути восстановления. Гипотеза: "
        "совместная процедура FIND–ERASE–PROVE уменьшает долю ложных сообщений о завершённом удалении "
        "по сравнению с локальными статусами и статическим аудитом, сохраняя меньшую стоимость действий, "
        "чем полная пересборка системы.",
    )
    add_section(
        doc,
        "Метод",
        "FIND строит карту данных и безопасными пробами уточняет ограниченное множество неочевидных "
        "путей восстановления. ERASE выбирает минимальный достаточный набор разрешённых действий; "
        "удаление влияния из модели (machine unlearning) является обязательным действием модельной ветки "
        "и сравнивается с полным переобучением. PROVE повторяет зарегистрированные будущие операции и выдаёт сертификат только "
        "при отсутствии остаточного или регенерируемого пути. При нехватке доказательств алгоритм "
        "возвращает UNVERIFIED, а не положительный результат.",
    )
    add_section(
        doc,
        "Основные результаты",
        "На 60 заранее зафиксированных случаях EraSeMap получил 0 ложных COMPLETE против 5 у полного аудита "
        "узлов и 45 у обычного статуса сервиса. Активный поиск использовал 7 проб против 13 у случайной "
        "стратегии и 49 у полного перебора. В 20 парных опытах с реальными локальными процессами выбранный "
        "план был в 17,64 раза быстрее полной пересборки и записал на 94,62% меньше байтов. PROVE обнаружил 30 из 30 скрытых рисков и после контрольных действий дал "
        "0 из 30 возвратов. Exact solvers совпали с exhaustive oracle в 3072 из 3072 и 16 384 из 16 384 "
        "конфигураций. Быстрые методы удаления влияния из Qwen не прошли все пороги; отрицательный результат "
        "сохранён, а полное переобучение оставлено безопасным запасным вариантом.",
    )
    add_section(
        doc,
        "Новизна и практическая значимость",
        "Новизна здесь не в отдельных известных идеях lineage, set cover или machine unlearning, "
        "а в одном строгом правиле решения: ни одна локальная квитанция не может самостоятельно объявить "
        "удаление завершённым. Подход применим к биометрии, банковским системам, электронным услугам, "
        "медицинским данным и другим системам с копиями, производными и моделями.",
    )
    add_section(
        doc,
        "Ограничения и вывод",
        "Гарантия действует только внутри зарегистрированной карты, verifier-каналов и временного окна; "
        "промышленное внедрение и независимый hidden challenge пока не заявляются. Эксперименты показывают, "
        "что проверять полный путь удаления безопаснее, чем доверять одной команде DELETE, а targeted план "
        "может быть существенно дешевле полной пересборки.",
    )
    doc.core_properties.title = f"Тезисы — {TITLE}"
    doc.core_properties.author = "Нұрланұлы Дулат"
    doc.save(path)


def build_bilingual_abstract(path: Path) -> None:
    if word_count(ABSTRACT_RU) > 250 or word_count(ABSTRACT_EN) > 250:
        raise ValueError("Abstract must not exceed 250 words in either language")
    doc = Document()
    configure_document(doc, True)
    add_document_title(doc, "АННОТАЦИЯ / ABSTRACT")
    add_identity(doc)
    add_section(doc, f"Аннотация ({word_count(ABSTRACT_RU)} слов)", ABSTRACT_RU)
    p = doc.add_paragraph()
    add_inline(
        p,
        "**Ключевые слова:** проверяемое удаление, machine unlearning, data lineage, temporal replay, "
        "биометрия, fail-closed, сертификат удаления.",
    )
    add_section(doc, f"Abstract ({word_count(ABSTRACT_EN)} words)", ABSTRACT_EN)
    p = doc.add_paragraph()
    add_inline(
        p,
        "**Keywords:** verifiable erasure, machine unlearning, data lineage, temporal replay, "
        "biometrics, fail-closed, deletion certificate.",
    )
    doc.core_properties.title = f"Аннотация / Abstract — {TITLE}"
    doc.core_properties.author = "Нұрланұлы Дулат"
    doc.save(path)


def build_supervisor_review(path: Path) -> None:
    doc = Document()
    configure_document(doc, True)
    add_document_title(doc, "ПРОЕКТ ОТЗЫВА НАУЧНОГО РУКОВОДИТЕЛЯ")
    add_identity(doc)
    warning = doc.add_paragraph()
    warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = warning.add_run("Документ действителен только после личной проверки, правок и подписи руководителя.")
    set_run_font(r, name="PT Sans", size=10, bold=True, color=BLUE)
    add_section(
        doc,
        "Актуальность",
        "Работа посвящена проверяемому удалению персональных данных из распределённых информационных "
        "систем и моделей машинного обучения. Проблема актуальна для биометрических, банковских, "
        "государственных и образовательных систем, где удаление основной записи не гарантирует исчезновение "
        "копий, производных и возможности будущего восстановления.",
    )
    add_section(
        doc,
        "Личный вклад автора",
        "Автор сформулировал исследовательский вопрос, разработал единый алгоритм FIND–ERASE–PROVE, "
        "реализовал программный прототип и воспроизводимые проверки, подготовил экспериментальные сценарии, "
        "проанализировал результаты и сохранил отрицательные результаты machine unlearning. В работе "
        "разделены подтверждённые результаты, ограничения и будущие направления исследования.",
    )
    add_section(
        doc,
        "Достоверность и самостоятельность",
        "Код, протоколы и сводные результаты опубликованы в репозитории. Точные "
        "решатели сопоставлены с полным перебором, а ключевые ограничения указаны явно. Перед подписанием "
        "руководителю нужно проверить, соответствует ли этот текст фактическому вкладу автора "
        "и результатам, представленным в работе.",
    )
    add_section(
        doc,
        "Недостатки и рекомендации",
        "Основные ограничения — отсутствие рабочего пилота и завершённого независимого скрытого "
        "эксперимента, а также отрицательный результат быстрых методов удаления влияния. Рекомендуется провести "
        "независимую проверку на заранее скрытых топологиях и, при наличии разрешения организации, пилот на "
        "реальной инфраструктуре без использования персональных данных пользователей.",
    )
    add_section(
        doc,
        "Заключение руководителя",
        "Работа имеет исследовательскую и практическую ценность, содержит программную реализацию, "
        "формальное описание алгоритма, измеряемые результаты и понятные границы применимости. После "
        "проверки фактического личного вклада и оформления документ может быть рекомендован к участию в секции "
        "«Информатика».",
    )
    doc.add_paragraph()
    for text in (
        "Научный руководитель: Смағұл Ерзат Айдынұлы",
        "Должность: учитель информатики",
        "Подпись: ____________________",
        "Дата: ____________________",
    ):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        r = p.add_run(text)
        set_run_font(r, name="PT Sans", size=10.5, color=DARK)
    doc.core_properties.title = f"Проект отзыва руководителя — {TITLE}"
    doc.core_properties.author = "Смағұл Ерзат Айдынұлы — подпись не собрана"
    doc.save(path)


def build_registration_sheet(path: Path) -> None:
    doc = Document()
    configure_document(doc, True)
    add_document_title(doc, "ДАННЫЕ ДЛЯ РЕГИСТРАЦИИ — НЕОФИЦИАЛЬНЫЙ ЧЕРНОВИК")
    rows = [
        ("Автор", "Нұрланұлы Дулат"),
        ("Имя латиницей", "Nurlanuly Dulat"),
        ("Класс", "9 «Б» / Grade 9B"),
        ("Организация", SCHOOL_RU),
        ("Город", "Алматы"),
        ("Формат", "Индивидуальный проект"),
        ("Язык", "Русский"),
        ("Направление", DIRECTION_RU),
        ("Секция", SECTION_RU),
        ("Название", TITLE),
        ("Руководитель", "Смағұл Ерзат Айдынұлы"),
        ("Должность", "Учитель информатики"),
        ("ИИН автора", "Заполнить только в закрытой официальной форме"),
        ("Телефон и e-mail", "Заполнить только в закрытой официальной форме"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (key, value) in enumerate(rows):
        table.cell(i, 0).text = key
        table.cell(i, 1).text = value
        for j, cell in enumerate(table.rows[i].cells):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, name="PT Sans", size=9.5, bold=(j == 0), color=MUTED if j == 0 else DARK)
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    add_inline(
        note,
        "**Важно:** это справочный лист, а не официальная заявка. Школа или координатор должны перенести "
        "данные в актуальный шаблон и проверить их по удостоверяющим документам.",
    )
    doc.core_properties.title = "Данные для регистрации EraSeMap — неофициальный черновик"
    doc.core_properties.author = "Нұрланұлы Дулат"
    doc.save(path)


def main() -> int:
    OUT.mkdir(parents=True, exist_ok=True)
    jobs = [
        (build_theses, OUT / "Nurlanuly_Dulat_EraSeMap_Tezisy_RU.docx"),
        (build_bilingual_abstract, OUT / "Nurlanuly_Dulat_EraSeMap_Annotatsiya_RU_EN.docx"),
        (build_supervisor_review, OUT / "Nurlanuly_Dulat_EraSeMap_Otzyv_Rukovoditelya_DRAFT.docx"),
        (build_registration_sheet, OUT / "Nurlanuly_Dulat_EraSeMap_Registration_Data_DRAFT.docx"),
    ]
    for builder, path in jobs:
        builder(path)
        print(path)
    print(f"RU abstract words: {word_count(ABSTRACT_RU)}")
    print(f"EN abstract words: {word_count(ABSTRACT_EN)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
