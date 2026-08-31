#!/usr/bin/env python3
# ruff: noqa: E501, RUF001
"""Build synchronized Russian and English EraSeMap paper deliverables."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[2]
PAPER_DIR = ROOT / "competition" / "paper"
ASSET_DIR = PAPER_DIR / "assets"
BUILD_DIR = PAPER_DIR / "build"

BLUE = "173F38"
MID_BLUE = "2C6F63"
DARK = "173F38"
MUTED = "5B7069"
PALE = "E8F2ED"
GRID = "CBDAD3"
WHITE = "FFFFFF"
RED = "C64F3A"
GREEN = "078C82"
SOFT_BLUE = "F4F8F6"
SOFT_GREEN = "EEF7F3"
PAPER = "FFFDF8"

AUTHOR_RU = "Нұрланұлы Дулат, 9 «Б» класс"
AUTHOR_EN = "Nurlanuly Dulat, Grade 9B"
SCHOOL_RU = "КГУ «Специализированный лицей-интернат „Білім-инновация“» Управления образования города Алматы"
SCHOOL_EN = "Specialized Bilim-Innovation Lyceum-Boarding School of the Almaty City Education Department"
SUPERVISOR_RU = "Смағұл Ерзат Айдынұлы, учитель информатики"
SUPERVISOR_EN = "Smagul Yerzat Aidynuly, Computer Science Teacher"
DIRECTION_RU = "Направление II — математическое моделирование экономических и социальных процессов"
SECTION_RU = "Секция: информатика"
DIRECTION_EN = "Direction II — mathematical modelling of economic and social processes"
SECTION_EN = "Section: Computer Science"

# narrative_proposal preset + named A4 academic-submission override.
PAGE_WIDTH_DXA = 11907
PAGE_HEIGHT_DXA = 16840
MARGIN_DXA = 1440
TABLE_WIDTH_DXA = PAGE_WIDTH_DXA - 2 * MARGIN_DXA
TABLE_INDENT_DXA = 120


PT_SANS = "/System/Library/Fonts/Supplemental/PTSans.ttc"
PT_SERIF = "/System/Library/Fonts/Supplemental/PTSerif.ttc"


def figure_font(size: int, *, bold: bool = False, italic: bool = False) -> ImageFont.FreeTypeFont:
    """Use the same Cyrillic-safe type family in figures and the manuscript."""
    if not Path(PT_SANS).exists():
        raise FileNotFoundError("PT Sans is required for the paper figures")
    index = 7 if bold else 1 if italic else 0
    return ImageFont.truetype(PT_SANS, size, index=index)


def center_text(
    d: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font: ImageFont.FreeTypeFont,
    fill: str,
) -> None:
    left, top, right, bottom = box
    bounds = d.multiline_textbbox((0, 0), text, font=font, spacing=8, align="center")
    d.multiline_text(
        ((left + right - (bounds[2] - bounds[0])) / 2, (top + bottom - (bounds[3] - bounds[1])) / 2),
        text,
        font=font,
        fill=fill,
        spacing=8,
        align="center",
    )


def make_system_figure(path: Path, ru: bool) -> None:
    img = Image.new("RGB", (2200, 1050), f"#{PAPER}")
    d = ImageDraw.Draw(img)
    title = figure_font(54, bold=True)
    subtitle = figure_font(29)
    stage = figure_font(34, bold=True)
    small = figure_font(25)
    d.text(
        (110, 72),
        "Один алгоритм EraSeMap: три шага"
        if ru
        else "One EraSeMap algorithm: three steps",
        fill=f"#{BLUE}",
        font=title,
    )
    d.text(
        (112, 145),
        "Один verdict только после обязательных доказательств всех каналов."
        if ru
        else "One verdict only after evidence from every required channel.",
        fill=f"#{MUTED}",
        font=subtitle,
    )
    request = (110, 310, 390, 570)
    d.rounded_rectangle(request, radius=28, fill="#FFFFFF", outline=f"#{GRID}", width=3)
    center_text(d, request, "Запрос\nна удаление" if ru else "Deletion\nrequest", stage, f"#{DARK}")

    stages = [
        (480, "1", "FIND", "Копии · модель\nскрытые пути" if ru else "Copies · model\nhidden paths"),
        (980, "2", "ERASE", "Физическое удаление\n+ machine unlearning" if ru else "Physical deletion\n+ machine unlearning"),
        (1480, "3", "PROVE", "Temporal replay\n+ сертификат" if ru else "Temporal replay\n+ certificate"),
    ]
    for x, number, name, description in stages:
        rect = (x, 270, x + 400, 610)
        d.rounded_rectangle(rect, radius=30, fill="#FFFFFF", outline=f"#{MID_BLUE}", width=4)
        d.ellipse((x + 32, 300, x + 102, 370), fill=f"#{BLUE}")
        center_text(d, (x + 32, 300, x + 102, 370), number, figure_font(28, bold=True), "#FFFFFF")
        d.text((x + 125, 314), name, font=stage, fill=f"#{BLUE}")
        center_text(d, (x + 40, 400, x + 360, 560), description, small, f"#{DARK}")
    for x in (410, 910, 1410):
        d.line((x, 440, x + 65, 440), fill=f"#{MID_BLUE}", width=8)
        d.polygon([(x + 80, 440), (x + 53, 423), (x + 53, 457)], fill=f"#{MID_BLUE}")

    d.text((110, 700), "Fail-closed результат" if ru else "Fail-closed output", font=stage, fill=f"#{BLUE}")
    outputs = [
        (110, "COMPLETE\nWITHIN ENVELOPE", "#E8F4EC", GREEN),
        (780, "INCOMPLETE", "#FFF6E6", "A06100"),
        (1350, "UNVERIFIED", "#F1F5F9", MUTED),
    ]
    for x, label, fill, outline in outputs:
        rect = (x, 780, x + 470, 940)
        d.rounded_rectangle(rect, radius=22, fill=fill, outline=f"#{outline}", width=3)
        center_text(d, rect, label, figure_font(29, bold=True), f"#{outline}")
    img.save(path, dpi=(220, 220))


def make_result_figure(path: Path, ru: bool) -> None:
    img = Image.new("RGB", (2200, 1680), f"#{PAPER}")
    d = ImageDraw.Draw(img)
    title = figure_font(54, bold=True)
    panel_title = figure_font(30, bold=True)
    label = figure_font(25)
    value = figure_font(42, bold=True)
    d.text((110, 70), "Результаты: сравнения с baseline" if ru else "Results: comparisons with baselines", fill=f"#{BLUE}", font=title)
    d.text(
        (112, 148),
        "Каждая панель — отдельный зафиксированный эксперимент; объединённый общий score не заявляется."
        if ru
        else "Each panel is a separate frozen experiment; no pooled overall score is claimed.",
        fill=f"#{MUTED}",
        font=label,
    )
    panels = [(110, 245, 1050, 870), (1150, 245, 2090, 870), (110, 965, 1050, 1570), (1150, 965, 2090, 1570)]
    for panel in panels:
        d.rounded_rectangle(panel, radius=26, fill="#FFFFFF", outline=f"#{GRID}", width=3)

    # A. False-complete rate: a direct safety comparison.
    left, top, right, _ = panels[0]
    d.text((left + 45, top + 40), "A. Безопасность verdict" if ru else "A. Verdict safety", font=panel_title, fill=f"#{DARK}")
    d.text((left + 45, top + 88), "Ложный COMPLETE — ниже лучше (stress test, n=75)" if ru else "False COMPLETE — lower is better (stress test, n=75)", font=label, fill=f"#{MUTED}")
    axis_x, axis_y, axis_width = left + 255, top + 235, 560
    for tick in (0, 25, 50, 75, 100):
        x = axis_x + int(axis_width * tick / 100)
        d.line((x, axis_y - 12, x, axis_y + 300), fill="#E2E8F0", width=2)
        d.text((x - 12, axis_y + 318), str(tick), fill=f"#{MUTED}", font=figure_font(20))
    for idx, (name, percent, numerator, color) in enumerate((("EraSeMap", 0, "0/75", GREEN), ("Typed-node", 100, "75/75", RED))):
        y = axis_y + idx * 135
        d.text((left + 45, y + 12), name, fill=f"#{DARK}", font=label)
        d.rounded_rectangle((axis_x, y, axis_x + axis_width, y + 58), radius=12, fill="#EEF2F6")
        if percent:
            d.rounded_rectangle((axis_x, y, axis_x + int(axis_width * percent / 100), y + 58), radius=12, fill=f"#{color}")
        else:
            d.ellipse((axis_x + 13, y + 13, axis_x + 45, y + 45), fill=f"#{color}")
        d.text((axis_x + axis_width + 26, y + 9), numerator, fill=f"#{color}", font=panel_title)
    d.text(
        (axis_x, axis_y + 376),
        "% false-positive verdicts" if not ru else "% неточных положительных verdict",
        fill=f"#{MUTED}",
        font=figure_font(20),
    )

    # B. Targeted ERASE cost compared with a full rebuild.
    left, top, right, _ = panels[1]
    d.text((left + 45, top + 40), "B. Реальное исполнение" if ru else "B. Real execution", font=panel_title, fill=f"#{DARK}")
    d.text((left + 45, top + 88), "20/20 COMPLETE; targeted ERASE против rebuild-all" if ru else "20/20 COMPLETE; targeted ERASE vs rebuild-all", font=label, fill=f"#{MUTED}")
    metrics = [
        ("Время записи" if ru else "Write time", "5.67%", "100%", "17.64× быстрее" if ru else "17.64× faster"),
        ("Записанные байты" if ru else "Written bytes", "5.38%", "100%", "94.62% меньше" if ru else "94.62% fewer"),
    ]
    for idx, (metric, targeted, rebuild, note) in enumerate(metrics):
        y = top + 210 + idx * 190
        d.text((left + 48, y), metric, font=panel_title, fill=f"#{DARK}")
        d.text((left + 48, y + 48), "EraSeMap", font=label, fill=f"#{MUTED}")
        d.rounded_rectangle((left + 245, y + 48, left + 770, y + 94), radius=10, fill="#E7EEF5")
        d.rounded_rectangle((left + 245, y + 48, left + 275, y + 94), radius=10, fill=f"#{BLUE}")
        d.text((left + 790, y + 51), targeted, font=panel_title, fill=f"#{BLUE}")
        d.text((left + 48, y + 112), "Rebuild-all", font=label, fill=f"#{MUTED}")
        d.rounded_rectangle((left + 245, y + 112, left + 770, y + 158), radius=10, fill="#D9E1E8")
        d.text((left + 790, y + 115), rebuild, font=panel_title, fill=f"#{DARK}")
        d.text((left + 245, y + 166), note, font=label, fill=f"#{GREEN}")

    # C. Temporal replay: three independently reported checks.
    left, top, right, bottom = panels[2]
    d.text((left + 45, top + 40), "C. Проверка во времени" if ru else "C. Temporal replay", font=panel_title, fill=f"#{DARK}")
    d.text((left + 45, top + 88), "Первый preregistered temporal run" if ru else "First preregistered temporal run", font=label, fill=f"#{MUTED}")
    temporal = [
        ("30/30", "рисков найдено" if ru else "risks detected"),
        ("10/10", "безопасных случаев" if ru else "guarded safe cases"),
        ("0/30", "повторов после controls" if ru else "recurrences after controls"),
    ]
    for idx, (number, caption) in enumerate(temporal):
        y = top + 185 + idx * 118
        d.ellipse((left + 55, y, left + 125, y + 70), fill=f"#{GREEN}")
        center_text(d, (left + 55, y, left + 125, y + 70), "✓", figure_font(34, bold=True), "#FFFFFF")
        d.text((left + 155, y - 2), number, font=value, fill=f"#{GREEN}")
        d.text((left + 370, y + 13), caption, font=label, fill=f"#{DARK}")
        if idx < 2:
            d.line((left + 55, y + 95, right - 45, y + 95), fill="#E2E8F0", width=2)
    d.rounded_rectangle((left + 45, bottom - 115, right - 45, bottom - 45), radius=14, fill=f"#{SOFT_GREEN}")
    center_text(d, (left + 45, bottom - 115, right - 45, bottom - 45), "Снимок без controls: 0/30; PROVE обнаружил будущие риски" if ru else "Snapshot without controls: 0/30; PROVE found future risks", figure_font(23, bold=True), f"#{GREEN}")

    # D. Probe-budget comparison on the bounded hidden graph.
    left, top, right, bottom = panels[3]
    d.text((left + 45, top + 40), "D. Активное обнаружение" if ru else "D. Active discovery", font=panel_title, fill=f"#{DARK}")
    d.text((left + 45, top + 88), "Число проб — ниже лучше (один hidden graph)" if ru else "Probe count — lower is better (one hidden graph)", font=label, fill=f"#{MUTED}")
    names = ["EraSeMap\nactive minimax", "Frozen\nrandom", "Exhaustive\nsearch"]
    vals = [7, 13, 49]
    colors = [BLUE, "94A3B8", "94A3B8"]
    base_y, max_height = bottom - 105, 260
    d.line((left + 75, base_y, right - 75, base_y), fill=f"#{GRID}", width=3)
    for idx, (name, metric, color) in enumerate(zip(names, vals, colors, strict=True)):
        x = left + 170 + idx * 260
        height = int(max_height * metric / 49)
        d.rounded_rectangle((x, base_y - height, x + 110, base_y), radius=12, fill=f"#{color}")
        d.text((x + 24, base_y - height - 58), str(metric), font=value, fill=f"#{color}")
        center_text(d, (x - 35, base_y + 22, x + 145, base_y + 98), name, figure_font(21, bold=(idx == 0)), f"#{DARK}")
    d.rounded_rectangle((left + 45, top + 470, right - 45, top + 545), radius=14, fill=f"#{SOFT_BLUE}")
    center_text(d, (left + 45, top + 470, right - 45, top + 545), "7 = active-minimax минимум; ничья с greedy, лучше random/exhaustive" if ru else "7 = active-minimax minimum; tied with greedy, better than random/exhaustive", figure_font(22, bold=True), f"#{BLUE}")
    img.save(path, dpi=(220, 220))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_borders(table, color: str = GRID, size: str = "6") -> None:
    """Apply light, consistent borders instead of Word's heavy Table Grid preset."""
    tbl_pr = table._tbl.tblPr
    for child in list(tbl_pr):
        if child.tag == qn("w:tblBorders"):
            tbl_pr.remove(child)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)
    tbl_pr.append(borders)


def set_run_font(run, name: str = "PT Serif", size: float | None = None, bold: bool | None = None,
                 italic: bool | None = None, color: str | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_inline(paragraph, text: str, *, default_bold: bool = False, default_italic: bool = False) -> None:
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`|\*.*?\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, bold=default_bold, italic=default_italic)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, bold=True, italic=default_italic)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Menlo", size=9.5, color=BLUE)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, bold=default_bold, italic=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, bold=default_bold, italic=default_italic)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_run_font(run, name="PT Sans", size=9, color=MUTED)


def configure_document(doc: Document, ru: bool) -> None:
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)
    sec.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "PT Serif"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "PT Serif")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "PT Serif")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PT Serif")
    normal._element.rPr.rFonts.set(qn("w:cs"), "PT Serif")
    normal.font.size = Pt(11.3)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(0)
    pf.space_after = Pt(7)
    pf.line_spacing = 1.23

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, MID_BLUE, 12, 6),
        "Heading 3": (12, BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = "PT Sans"
        style._element.rPr.rFonts.set(qn("w:ascii"), "PT Sans")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "PT Sans")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PT Sans")
        style._element.rPr.rFonts.set(qn("w:cs"), "PT Sans")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "PT Serif"
        style._element.rPr.rFonts.set(qn("w:ascii"), "PT Serif")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "PT Serif")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PT Serif")
        style._element.rPr.rFonts.set(qn("w:cs"), "PT Serif")
        style.font.size = Pt(11.3)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    for name in ("Formula", "Figure Caption", "Callout", "Code Block", "Numbered Item"):
        if name not in [style.name for style in doc.styles]:
            doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    formula = doc.styles["Formula"]
    formula.font.name = "Cambria Math"
    formula.font.size = Pt(11.5)
    formula.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula.paragraph_format.space_before = Pt(4)
    formula.paragraph_format.space_after = Pt(8)
    formula.paragraph_format.keep_together = True
    caption = doc.styles["Figure Caption"]
    caption.font.name = "PT Sans"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "PT Sans")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "PT Sans")
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "PT Sans")
    caption._element.rPr.rFonts.set(qn("w:cs"), "PT Sans")
    caption.font.size = Pt(9.2)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)
    callout = doc.styles["Callout"]
    callout.font.name = "PT Serif"
    callout._element.rPr.rFonts.set(qn("w:ascii"), "PT Serif")
    callout._element.rPr.rFonts.set(qn("w:hAnsi"), "PT Serif")
    callout._element.rPr.rFonts.set(qn("w:eastAsia"), "PT Serif")
    callout._element.rPr.rFonts.set(qn("w:cs"), "PT Serif")
    callout.font.size = Pt(11.3)
    callout.font.italic = True
    callout.font.color.rgb = RGBColor.from_string(BLUE)
    callout.paragraph_format.left_indent = Inches(0.25)
    callout.paragraph_format.right_indent = Inches(0.25)
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(10)
    callout.paragraph_format.line_spacing = 1.18
    code = doc.styles["Code Block"]
    code.font.name = "Menlo"
    code.font.size = Pt(8.5)
    code.paragraph_format.left_indent = Inches(0.25)
    code.paragraph_format.space_after = Pt(2)
    code.paragraph_format.line_spacing = 1.0
    numbered = doc.styles["Numbered Item"]
    numbered.font.name = "PT Serif"
    numbered._element.rPr.rFonts.set(qn("w:ascii"), "PT Serif")
    numbered._element.rPr.rFonts.set(qn("w:hAnsi"), "PT Serif")
    numbered._element.rPr.rFonts.set(qn("w:eastAsia"), "PT Serif")
    numbered._element.rPr.rFonts.set(qn("w:cs"), "PT Serif")
    numbered.font.size = Pt(11.3)
    numbered.paragraph_format.left_indent = Inches(0.375)
    numbered.paragraph_format.first_line_indent = Inches(-0.194)
    numbered.paragraph_format.space_after = Pt(4)
    numbered.paragraph_format.line_spacing = 1.208

    header = sec.header.paragraphs[0]
    header.text = "EraSeMap | " + ("Научная работа" if ru else "Research paper")
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(header.runs[0], name="PT Sans", size=9, bold=True, color=MUTED)
    footer = sec.footer.paragraphs[0]
    add_page_field(footer)


def add_cover(doc: Document, title: str, ru: bool) -> None:
    for _ in range(5):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("НАУЧНО-ИССЛЕДОВАТЕЛЬСКАЯ РАБОТА" if ru else "SCIENTIFIC RESEARCH PAPER")
    set_run_font(run, name="PT Sans", size=11, bold=True, color=MID_BLUE)
    kicker.paragraph_format.space_after = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(title)
    set_run_font(r, name="PT Sans", size=27, bold=True, color=BLUE)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(52)
    r = sub.add_run("A unified FIND–ERASE–PROVE algorithm" if not ru else "Единый алгоритм FIND–ERASE–PROVE")
    set_run_font(r, name="PT Sans", size=14, italic=True, color=MUTED)

    fields = [
        ("Автор" if ru else "Author", AUTHOR_RU if ru else AUTHOR_EN),
        ("Организация" if ru else "Affiliation", SCHOOL_RU if ru else SCHOOL_EN),
        ("Направление" if ru else "Direction", DIRECTION_RU if ru else DIRECTION_EN),
        ("Секция" if ru else "Section", SECTION_RU if ru else SECTION_EN),
        ("Научный руководитель" if ru else "Supervisor", SUPERVISOR_RU if ru else SUPERVISOR_EN),
    ]
    table = doc.add_table(rows=len(fields), cols=2)
    widths = [2450, TABLE_WIDTH_DXA - 2450]
    set_table_geometry(table, widths)
    for i, (key, value) in enumerate(fields):
        table.cell(i, 0).text = key
        table.cell(i, 1).text = value
        for j, cell in enumerate(table.rows[i].cells):
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(4)
                for run in paragraph.runs:
                    set_run_font(run, name="PT Sans", size=11, bold=(j == 0), color=MUTED if j == 0 else DARK)
            set_cell_shading(cell, WHITE)
    # Remove table borders on the cover.
    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        tbl_borders.append(el)
    table._tbl.tblPr.append(tbl_borders)
    # The metadata grid is a layout table; mark its first row so assistive
    # technology does not report an unlabeled table structure.
    set_repeat_table_header(table.rows[0])

    # Keep the cover on one A4 page even when a Russian metadata label wraps.
    doc.add_paragraph()
    year = doc.add_paragraph()
    year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = year.add_run("Алматы, 2026" if ru else "Almaty, 2026")
    set_run_font(r, name="PT Sans", size=11, bold=True, color=MUTED)
    doc.add_page_break()


def add_contents(doc: Document, ru: bool) -> None:
    heading = doc.add_paragraph("Оглавление" if ru else "Contents", style="Heading 1")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    items = (
        [
            ("Аннотация", 3),
            ("Abstract", 3),
            ("1. Введение", 4),
            ("2. Предшествующие работы и граница новизны", 4),
            ("3. Модель системы", 5),
            ("4. Единый алгоритм EraSeMap", 6),
            ("5. Формальные свойства", 7),
            ("6. Реализация", 7),
            ("7. Методика экспериментов", 8),
            ("8. Результаты", 8),
            ("9. Обсуждение", 9),
            ("10. Ограничения и угрозы валидности", 10),
            ("11. Этика и ответственное применение", 10),
            ("12. Воспроизводимость", 10),
            ("13. Заключение", 11),
            ("Список литературы", 11),
            ("Приложения", 12),
        ]
        if ru
        else [
            ("Abstract", 3),
            ("1. Introduction", 4),
            ("2. Related work and novelty boundary", 4),
            ("3. System model", 5),
            ("4. The unified EraSeMap algorithm", 6),
            ("5. Formal properties", 7),
            ("6. Implementation", 7),
            ("7. Experimental methodology", 8),
            ("8. Results", 8),
            ("9. Discussion", 9),
            ("10. Limitations and threats to validity", 10),
            ("11. Ethics and responsible use", 10),
            ("12. Reproducibility", 10),
            ("13. Conclusion", 11),
            ("References", 11),
            ("Appendices", 12),
        ]
    )
    for item, page in items:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.35)
        paragraph.paragraph_format.space_after = Pt(3.5)
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.1), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )
        run = paragraph.add_run(f"{item}\t{page}")
        set_run_font(run, name="PT Sans", size=10.5, color=DARK)
    doc.add_page_break()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            rows.append(cells)
        i += 1
    return rows, i


def table_widths(cols: int) -> list[int]:
    if cols == 2:
        return [3000, TABLE_WIDTH_DXA - 3000]
    if cols == 3:
        return [2300, 2450, TABLE_WIDTH_DXA - 4750]
    if cols == 4:
        return [1900, 1750, 1750, TABLE_WIDTH_DXA - 5400]
    base = TABLE_WIDTH_DXA // cols
    return [base] * (cols - 1) + [TABLE_WIDTH_DXA - base * (cols - 1)]


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    set_table_geometry(table, table_widths(cols))
    set_table_borders(table)
    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.1
            add_inline(p, text, default_bold=(i == 0))
            for run in p.runs:
                set_run_font(run, name="PT Sans", size=9.3, bold=(i == 0 or run.bold))
            if i == 0:
                set_cell_shading(cell, PALE)
            elif i % 2 == 0:
                set_cell_shading(cell, SOFT_BLUE)
    set_repeat_table_header(table.rows[0])
    for row in table.rows:
        prevent_row_split(row)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_figure(doc: Document, path: Path, caption: str, alt_text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    shape = p.add_run().add_picture(str(path), width=Inches(6.1))
    shape._inline.docPr.set("descr", alt_text)
    shape._inline.docPr.set("title", caption.split(".", 1)[0])
    cap = doc.add_paragraph(caption, style="Figure Caption")
    cap.paragraph_format.keep_with_next = False


def build_from_markdown(source: Path, output: Path, ru: bool) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    title = lines[0][2:].strip()
    doc = Document()
    configure_document(doc, ru)
    add_cover(doc, title, ru)
    add_contents(doc, ru)

    system_fig = ASSET_DIR / ("system-flow-ru.png" if ru else "system-flow-en.png")
    result_fig = ASSET_DIR / ("results-ru.png" if ru else "results-en.png")
    inserted_system = False
    inserted_results = False
    in_code = False
    skip_metadata = True

    def structural(text: str) -> bool:
        return bool(
            text.startswith(("#", "|", "> ", "- ", "```"))
            or re.match(r"^\d+\. ", text)
        )

    i = 1
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if skip_metadata:
            if line.startswith("## "):
                skip_metadata = False
            else:
                i += 1
                continue
        if line.startswith("```"):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            p = doc.add_paragraph(style="Code Block")
            p.add_run(raw)
            i += 1
            continue
        if not line:
            i += 1
            continue
        if line.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        if line.startswith("## "):
            heading = line[3:]
            doc.add_paragraph(heading, style="Heading 1")
            if heading.startswith("2.") and not inserted_system:
                add_figure(
                    doc,
                    system_fig,
                    ("Рисунок 1. Один алгоритм EraSeMap: FIND, ERASE и PROVE. Составлено автором по протоколу EraSeMap." if ru else "Figure 1. One EraSeMap algorithm: FIND, ERASE, and PROVE. Author-generated from the EraSeMap protocol."),
                    ("Схема: один запрос удаления проходит три обязательных шага FIND, ERASE и PROVE и заканчивается COMPLETE_WITHIN_ENVELOPE, INCOMPLETE или UNVERIFIED." if ru else "Flow diagram: one deletion request passes through the three mandatory FIND, ERASE, and PROVE steps and ends in COMPLETE_WITHIN_ENVELOPE, INCOMPLETE, or UNVERIFIED."),
                )
                inserted_system = True
            if (heading in ("9. Discussion", "9. Обсуждение")) and not inserted_results:
                add_figure(
                    doc,
                    result_fig,
                    ("Рисунок 2. Четыре измеренных свойства единого алгоритма EraSeMap: безопасность, эффективность, временная проверка и активное обнаружение. Составлено автором по зафиксированным результатам." if ru else "Figure 2. Four measured properties of the unified EraSeMap algorithm: safety, efficiency, temporal verification, and active discovery. Author-generated from the frozen results."),
                    ("Диаграмма: EraSeMap даёт 0 из 75 ложных COMPLETE против 75 из 75 у typed-node; работает в 17,64 раза быстрее rebuild-all и записывает на 94,62 процента меньше байтов; временной этап обнаруживает 30 из 30 рисков без повторов; активное обнаружение использует 7 probes против 13 random и 49 exhaustive." if ru else "Results chart: EraSeMap has 0 of 75 false COMPLETE verdicts versus 75 of 75 for typed-node; it is 17.64 times faster than rebuild-all with 94.62 percent fewer written bytes; its temporal stage detects 30 of 30 risks without recurrence; active discovery uses 7 probes versus 13 random and 49 exhaustive."),
                )
                inserted_results = True
            i += 1
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:], style="Heading 2")
            i += 1
            continue
        if line.startswith("#### "):
            doc.add_paragraph(line[5:], style="Heading 3")
            i += 1
            continue
        if line.startswith("> "):
            parts = [line[2:]]
            i += 1
            while i < len(lines) and lines[i].strip().startswith("> "):
                parts.append(lines[i].strip()[2:])
                i += 1
            p = doc.add_paragraph(style="Callout")
            add_inline(p, " ".join(parts))
            continue
        numbered_match = re.match(r"^(\d+)\. (.*)", line)
        if numbered_match:
            parts = [numbered_match.group(2)]
            i += 1
            while i < len(lines):
                continuation = lines[i].strip()
                if not continuation or structural(continuation):
                    break
                parts.append(continuation)
                i += 1
            p = doc.add_paragraph(style="Numbered Item")
            prefix = p.add_run(f"{numbered_match.group(1)}. ")
            set_run_font(prefix)
            add_inline(p, " ".join(parts))
            continue
        if line.startswith("- "):
            parts = [line[2:]]
            i += 1
            while i < len(lines):
                continuation = lines[i].strip()
                if not continuation or structural(continuation):
                    break
                parts.append(continuation)
                i += 1
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, " ".join(parts))
            continue
        if line.startswith("**") and line.endswith("**") and len(line) > 4:
            p = doc.add_paragraph(style="Formula")
            text = line[2:-2]
            run = p.add_run(text)
            set_run_font(run, name="Cambria Math", size=11.5)
            i += 1
            continue
        parts = [line]
        i += 1
        while i < len(lines):
            continuation = lines[i].strip()
            if not continuation or structural(continuation):
                break
            parts.append(continuation)
            i += 1
        p = doc.add_paragraph()
        p.paragraph_format.widow_control = True
        add_inline(p, " ".join(parts))

    core = doc.core_properties
    core.title = title
    core.subject = "Unified proof-carrying and regeneration-safe biometric erasure auditing"
    core.author = "Нұрланұлы Дулат" if ru else "Nurlanuly Dulat"
    core.keywords = "biometric erasure, machine unlearning, data lineage, verifiable deletion, regeneration-safe erasure"
    core.comments = "Design preset: narrative_proposal; named override: A4 academic submission geometry."
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    BUILD_DIR.mkdir(parents=True, exist_ok=True)
    make_system_figure(ASSET_DIR / "system-flow-en.png", False)
    make_system_figure(ASSET_DIR / "system-flow-ru.png", True)
    make_result_figure(ASSET_DIR / "results-en.png", False)
    make_result_figure(ASSET_DIR / "results-ru.png", True)
    jobs = [
        (PAPER_DIR / "EraSeMap_scientific_paper_EN.md", PAPER_DIR / "EraSeMap_scientific_paper_EN.docx", False),
        (PAPER_DIR / "EraSeMap_scientific_paper_RU.md", PAPER_DIR / "EraSeMap_scientific_paper_RU.docx", True),
    ]
    for source, output, ru in jobs:
        build_from_markdown(source, output, ru)
        print(output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
